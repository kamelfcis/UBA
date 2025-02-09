import docx

def create_docx_from_markdown(markdown_file, docx_file):
    # Create a new Document
    doc = docx.Document()
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8-sig') as file:
        content = file.readlines()
    
    # Add content to the DOCX file
    for line in content:
        if line.startswith('# '):  # Main heading
            doc.add_heading(line.strip('# ').strip(), level=1)
        elif line.startswith('## '):  # Subheading
            doc.add_heading(line.strip('## ').strip(), level=2)
        elif line.startswith('### '):  # Sub-subheading
            doc.add_heading(line.strip('### ').strip(), level=3)
        else:  # Regular text
            doc.add_paragraph(line.strip())
    
    # Save the document
    doc.save(docx_file)

# Generate the DOCX file
create_docx_from_markdown('documentation.md', 'documentation.docx')
